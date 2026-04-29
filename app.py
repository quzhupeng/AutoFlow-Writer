# -*- coding: utf-8 -*-
"""
企业流程架构设计智能体 - Streamlit主应用
基于 APQC 国际标准的端到端流程架构设计
"""

import os
import io
import streamlit as st
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from config import (DEEPSEEK_CONFIG, FULL_EXECUTION_PROMPT, GRANULARITY_RULES,
                    REASONING_LOGIC, SELF_CHECK_RULES, SYSTEM_PROMPT,
                    EXISTING_DOC_TEMPLATE, BPM_AUDIT_SYSTEM_PROMPT, BPM_AUDIT_USER_PROMPT)
from llm_client import DeepSeekClient

# ==================== 页面配置 ====================
st.set_page_config(
    page_title="流程架构设计与审核智能体",
    page_icon="🏗️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ==================== 工具函数 ====================

MAX_UPLOAD_SIZE_MB = 10
MAX_REFERENCE_CHARS = 12000

# 树状图美化样式
TREE_STYLES = """
<style>
.arch-tree {
    font-family: 'SF Mono', 'Consolas', 'Monaco', monospace;
    background: #ffffff;
    border: 1px solid #e0e0e0;
    border-radius: 8px;
    padding: 20px 25px;
    color: #333;
    line-height: 1.9;
    overflow-x: auto;
    box-shadow: 0 2px 8px rgba(0,0,0,0.06);
}
.arch-tree .l1 {
    font-size: 17px;
    font-weight: 700;
    color: #1a73e8;
    margin: 14px 0 8px 0;
    padding-bottom: 4px;
    border-bottom: 2px solid #e8f0fe;
}
.arch-tree .l2 {
    font-size: 15px;
    font-weight: 600;
    color: #137333;
    margin: 10px 0 6px 20px;
}
.arch-tree .l3 {
    font-size: 14px;
    color: #5f6368;
    margin: 6px 0 4px 40px;
}
.arch-tree .branch {
    color: #9aa0a6;
}
</style>
"""


def format_tree_display(tree_text: str) -> str:
    """将树状图文本转换为带样式的HTML"""
    lines = tree_text.strip().split('\n')
    html_parts = ['<div class="arch-tree">']

    for line in lines:
        if not line.strip():
            continue

        stripped = line.strip()
        indent = len(line) - len(line.lstrip())

        if stripped.startswith('L1：') or stripped.startswith('L1:'):
            html_parts.append(f'<div class="l1">{stripped}</div>')
        elif stripped.startswith('L2：') or stripped.startswith('L2:'):
            html_parts.append(f'<div class="l2">{stripped}</div>')
        elif stripped.startswith('L3：') or stripped.startswith('L3:'):
            html_parts.append(f'<div class="l3">{stripped}</div>')
        elif '├──' in stripped or '└──' in stripped:
            html_parts.append(f'<div class="branch" style="margin-left:{indent*2}px">{stripped}</div>')
        else:
            html_parts.append(f'<div style="margin-left:{indent*2}px">{stripped}</div>')

    html_parts.append('</div>')
    return '\n'.join(html_parts)


def display_message_with_tree(content: str):
    """智能显示消息内容，自动美化树状图"""
    has_tree = 'L1：' in content or 'L1:' in content or '├──' in content or '└──' in content

    if has_tree:
        lines = content.split('\n')
        tree_lines = []
        other_lines = []
        in_tree = False

        for line in lines:
            stripped = line.strip()
            if stripped.startswith('L1：') or stripped.startswith('L1:'):
                in_tree = True
                tree_lines.append(line)
            elif in_tree and (stripped.startswith('L2') or stripped.startswith('L3') or
                              '├──' in stripped or '└──' in stripped or not stripped):
                tree_lines.append(line)
            else:
                if in_tree and stripped and not stripped.startswith(('L2', 'L3', '├', '└', '│')):
                    in_tree = False
                if in_tree:
                    tree_lines.append(line)
                else:
                    other_lines.append(line)

        other_text = '\n'.join(other_lines).strip()
        if other_text:
            st.markdown(other_text)

        if tree_lines:
            tree_text = '\n'.join(tree_lines)
            st.markdown(format_tree_display(tree_text), unsafe_allow_html=True)
    else:
        st.markdown(content)

def extract_file_text(uploaded_file) -> str:
    """从上传的文件中提取文本内容"""
    if uploaded_file is None:
        return ""

    filename = uploaded_file.name.lower()

    try:
        content = uploaded_file.read()
        if not content:
            return ""

        if filename.endswith('.txt'):
            for encoding in ['utf-8', 'gbk', 'gb2312']:
                try:
                    return content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            return content.decode('utf-8', errors='ignore')

        elif filename.endswith('.docx'):
            doc = Document(io.BytesIO(content))
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            return '\n'.join(text_parts)

        elif filename.endswith('.pdf'):
            try:
                import PyPDF2
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(content))
                text_parts = []
                for page in pdf_reader.pages:
                    page_text = page.extract_text() or ""
                    if page_text.strip():
                        text_parts.append(page_text)
                return '\n'.join(text_parts)
            except ImportError:
                return "[错误: 需要安装PyPDF2库来处理PDF文件]"

        else:
            return f"[不支持的文件格式: {filename}]"

    except Exception as e:
        return f"[文件解析错误: {str(e)}]"


def normalize_reference_doc(text: str, max_chars: int = MAX_REFERENCE_CHARS):
    """规范化参考文档内容，必要时做截断以控制上下文长度"""
    normalized = (text or "").strip()
    if not normalized or len(normalized) <= max_chars:
        return normalized, False

    truncated = normalized[:max_chars]
    notice = f"\n\n[参考文档超长，已截断，仅保留前{max_chars}字符，原始长度{len(normalized)}字符]"
    return truncated + notice, True


def tree_to_docx(tree_text: str, department_name: str) -> Document:
    """将树状图文本转换为Word文档"""
    doc = Document()

    title_para = doc.add_heading(f"{department_name}端到端三级流程架构键盘图", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    from datetime import datetime
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_para.add_run(f"生成时间：{datetime.now().strftime('%Y年%m月%d日')}")
    doc.add_paragraph()

    lines = tree_text.split('\n')
    for line in lines:
        if not line.strip():
            continue

        if line.strip().startswith('L1：') or line.strip().startswith('L1:'):
            para = doc.add_heading(line.strip(), level=1)
        elif line.strip().startswith('L2：') or line.strip().startswith('L2:'):
            para = doc.add_heading(line.strip(), level=2)
        elif line.strip().startswith('L3：') or line.strip().startswith('L3:'):
            para = doc.add_paragraph(line.strip())
            para.paragraph_format.left_indent = Inches(0.5)
        else:
            para = doc.add_paragraph(line)

    return doc


def audit_report_to_docx(report_text: str) -> Document:
    """将审核报告Markdown转换为Word文档"""
    doc = Document()

    title_para = doc.add_heading("流程说明书评审报告", level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    from datetime import datetime
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info_para.add_run(f"审核时间：{datetime.now().strftime('%Y年%m月%d日')}")
    doc.add_paragraph()

    lines = report_text.split('\n')
    for line in lines:
        if not line.strip():
            doc.add_paragraph()
            continue
        if line.startswith('## '):
            doc.add_heading(line.replace('## ', '').strip(), level=1)
        elif line.startswith('### '):
            doc.add_heading(line.replace('### ', '').strip(), level=2)
        elif line.startswith('| ') and not all(set(c.strip()) <= set('-:| ') for c in line.split('|') if c.strip()):
            cells = [c.strip() for c in line.split('|') if c.strip()]
            if cells:
                table = doc.add_table(rows=1, cols=len(cells))
                table.style = 'Table Grid'
                for i, cell_text in enumerate(cells):
                    cell = table.rows[0].cells[i]
                    cell.text = cell_text
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True
        elif line.startswith('- '):
            para = doc.add_paragraph(line[2:], style='List Bullet')
        else:
            doc.add_paragraph(line)

    return doc

# ==================== Session State 初始化 ====================

def init_session_state():
    """初始化会话状态"""
    # 流程架构设计
    if 'messages' not in st.session_state:
        st.session_state.messages = []
    if 'architecture_result' not in st.session_state:
        st.session_state.architecture_result = None
    if 'department_name' not in st.session_state:
        st.session_state.department_name = ""
    if 'business_areas' not in st.session_state:
        st.session_state.business_areas = ""
    if 'existing_doc_text' not in st.session_state:
        st.session_state.existing_doc_text = None

    # 流程审核
    if 'audit_messages' not in st.session_state:
        st.session_state.audit_messages = []
    if 'audit_result' not in st.session_state:
        st.session_state.audit_result = None
    if 'audit_doc_text' not in st.session_state:
        st.session_state.audit_doc_text = ""


# ==================== 流程架构设计模块 ====================

def render_design_page(api_key, base_url, model):
    """流程架构设计页面"""
    st.header("📋 流程架构设计")

    col1, col2 = st.columns(2)
    with col1:
        department_name = st.text_input(
            "部门/业务模块全称",
            value=st.session_state.department_name,
            placeholder="例如：人力资源部、供应链管理部",
            key="design_dept"
        )

    with col2:
        business_areas = st.text_area(
            "核心细分业务板块",
            value=st.session_state.business_areas,
            placeholder="例如：招聘、培训、绩效、薪酬",
            height=80,
            key="design_areas"
        )

    st.subheader("📄 上传现有流程文档（可选）")
    uploaded_file = st.file_uploader("上传参考文档", type=['txt', 'docx', 'pdf'], key="design_upload")

    existing_doc_text = None
    if uploaded_file:
        file_size = getattr(uploaded_file, "size", 0)
        if file_size and file_size > MAX_UPLOAD_SIZE_MB * 1024 * 1024:
            st.error(f"文件过大：{file_size / (1024 * 1024):.1f}MB，最大支持 {MAX_UPLOAD_SIZE_MB}MB")
        else:
            with st.spinner("解析中..."):
                parsed_text = extract_file_text(uploaded_file)

            if parsed_text.startswith("["):
                st.warning(parsed_text)
            else:
                existing_doc_text, is_truncated = normalize_reference_doc(parsed_text)
                if existing_doc_text:
                    st.success(f"✓ 已解析: {uploaded_file.name}")
                    if is_truncated:
                        st.info(f"参考文档内容较长，已自动截断到前 {MAX_REFERENCE_CHARS} 字符。")
                else:
                    st.warning("文件已上传，但未提取到有效文本。")

    st.markdown("---")

    col_btn1, col_btn2, col_btn3 = st.columns([1, 1, 1])
    with col_btn2:
        start_button = st.button(
            "🚀 开始设计",
            type="primary",
            use_container_width=True,
            disabled=not (api_key and department_name.strip() and business_areas.strip()),
            key="start_design"
        )

    # ==================== 生成初始架构 ====================
    if start_button:
        st.session_state.messages = []
        st.session_state.architecture_result = None
        st.session_state.department_name = department_name
        st.session_state.business_areas = business_areas
        st.session_state.existing_doc_text = existing_doc_text

        progress_bar = st.progress(0)
        status_text = st.empty()

        try:
            status_text.text("正在连接 API...")
            progress_bar.progress(10)

            client = DeepSeekClient(api_key=api_key, base_url=base_url, model=model)

            status_text.text("🏗️ 正在构建三级流程架构...")
            progress_bar.progress(50)

            doc_section = ""
            if existing_doc_text and existing_doc_text.strip():
                doc_section = EXISTING_DOC_TEMPLATE.format(doc_content=existing_doc_text.strip())

            initial_prompt = FULL_EXECUTION_PROMPT.format(
                department_name=department_name,
                business_areas=business_areas,
                existing_process_doc_section=doc_section,
                granularity_rules=GRANULARITY_RULES,
                reasoning_logic=REASONING_LOGIC,
                self_check_rules=SELF_CHECK_RULES
            )

            st.session_state.messages = [
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": initial_prompt}
            ]

            result = client.chat(st.session_state.messages)
            if not result or not result.strip():
                raise RuntimeError("模型返回为空，请重试。")

            st.session_state.messages.append({"role": "assistant", "content": result})
            st.session_state.architecture_result = result

            progress_bar.progress(100)
            status_text.text("✓ 架构设计完成！可以继续对话优化")

        except Exception as e:
            progress_bar.empty()
            status_text.empty()
            st.error(f"❌ 执行失败: {str(e)}")
            return

    # ==================== 显示当前架构结果 ====================
    if st.session_state.architecture_result:
        st.markdown("---")
        st.header(f"📊 {st.session_state.department_name}端到端三级流程架构键盘图")

        st.markdown(TREE_STYLES, unsafe_allow_html=True)
        st.markdown(format_tree_display(st.session_state.architecture_result), unsafe_allow_html=True)

        with st.expander("📄 查看原始文本"):
            st.code(st.session_state.architecture_result, language=None)

        st.markdown("---")
        col_dl1, col_dl2, col_dl3 = st.columns([1, 1, 1])
        with col_dl2:
            doc = tree_to_docx(
                st.session_state.architecture_result,
                st.session_state.department_name or '企业'
            )
            doc_bytes = io.BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)

            st.download_button(
                label="📥 下载Word文档",
                data=doc_bytes,
                file_name=f"{st.session_state.department_name or '流程架构'}_架构键盘图.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                key="download_design"
            )

    # ==================== 多轮对话区域 ====================
    if st.session_state.messages:
        st.markdown("---")
        st.header("💬 继续对话优化架构")

        chat_messages = st.session_state.messages[3:] if len(st.session_state.messages) > 3 else []

        for msg in chat_messages:
            with st.chat_message(msg["role"]):
                display_message_with_tree(msg["content"])

        if prompt := st.chat_input("输入修改意见，如：'招聘流程L2需要拆分为社会招聘和校园招聘'", key="design_chat"):
            st.session_state.messages.append({"role": "user", "content": prompt})

            with st.chat_message("user"):
                st.markdown(prompt)

            with st.chat_message("assistant"):
                with st.spinner("思考中..."):
                    try:
                        client = DeepSeekClient(api_key=api_key, base_url=base_url, model=model)
                        response = client.chat(st.session_state.messages)
                        display_message_with_tree(response)

                        st.session_state.messages.append({"role": "assistant", "content": response})

                        if "L1：" in response or "L1:" in response:
                            st.session_state.architecture_result = response
                            st.rerun()

                    except Exception as e:
                        st.error(f"对话失败: {str(e)}")


# ==================== 流程审核模块 ====================

def render_audit_page(api_key, base_url, model):
    """流程审核页面"""
    st.header("🔍 流程说明书审核")
    st.markdown("**基于 16 项检查点的自动化评审，对标顶级咨询公司交付级别**")

    st.markdown("---")

    # 上传区域
    st.subheader("📄 上传流程说明书")

    # 文本粘贴
    manual_text = st.text_area(
        "直接粘贴流程文档内容",
        value=st.session_state.audit_doc_text,
        height=200,
        placeholder="将流程说明书全文粘贴到此处...",
        key="audit_text_input"
    )
    if manual_text != st.session_state.audit_doc_text:
        st.session_state.audit_doc_text = manual_text

    st.markdown("**或上传文件：**")
    uploaded_file = st.file_uploader("上传流程说明书", type=['txt', 'docx', 'pdf'], key="audit_upload")

    if uploaded_file:
        file_size = getattr(uploaded_file, "size", 0)
        if file_size and file_size > MAX_UPLOAD_SIZE_MB * 1024 * 1024:
            st.error(f"文件过大：{file_size / (1024 * 1024):.1f}MB，最大支持 {MAX_UPLOAD_SIZE_MB}MB")
        else:
            with st.spinner("解析中..."):
                parsed_text = extract_file_text(uploaded_file)
            if not parsed_text.startswith("["):
                st.session_state.audit_doc_text = parsed_text
                st.success(f"✓ 已解析: {uploaded_file.name}")
            else:
                st.warning(parsed_text)

    # 显示已输入内容长度
    if st.session_state.audit_doc_text:
        st.info(f"已输入 {len(st.session_state.audit_doc_text)} 字符")

    # 审核维度说明
    with st.expander("📊 审核维度说明"):
        st.markdown("""
        | 维度 | 检查项 | 说明 |
        | :--- | :--- | :--- |
        | 基础信息 | 4项 | 流程目的、适用范围、角色定义、术语解释 |
        | 业务逻辑 | 4项 | 战略对齐、PDCA闭环、起止边界、输入输出 |
        | 执行规范 | 4项 | 5W2H完整性、IT化、活动切分、咬合逻辑 |
        | 风控资源 | 4项 | 附件实操性、知识沉淀、风险管理、KPI设置 |
        """)

    st.markdown("---")

    # 开始审核
    col_btn1, col_btn2, col_btn3 = st.columns([1, 1, 1])
    with col_btn2:
        audit_button = st.button(
            "🔍 开始审核",
            type="primary",
            use_container_width=True,
            disabled=not (api_key and st.session_state.audit_doc_text.strip()),
            key="start_audit"
        )

    # ==================== 执行审核 ====================
    if audit_button:
        st.session_state.audit_messages = []
        st.session_state.audit_result = None

        progress_bar = st.progress(0)
        status_text = st.empty()

        try:
            status_text.text("正在进行流程审核...")
            progress_bar.progress(10)

            client = DeepSeekClient(api_key=api_key, base_url=base_url, model=model)

            status_text.text("🔍 AI 正在基于 16 项检查点逐项审计...")
            progress_bar.progress(50)

            user_prompt = BPM_AUDIT_USER_PROMPT.format(
                document_content=st.session_state.audit_doc_text.strip()
            )

            st.session_state.audit_messages = [
                {"role": "system", "content": BPM_AUDIT_SYSTEM_PROMPT},
                {"role": "user", "content": user_prompt}
            ]

            result = client.chat(st.session_state.audit_messages)
            if not result or not result.strip():
                raise RuntimeError("模型返回为空，请重试。")

            st.session_state.audit_messages.append({"role": "assistant", "content": result})
            st.session_state.audit_result = result

            progress_bar.progress(100)
            status_text.text("✓ 流程审核完成！")

        except Exception as e:
            progress_bar.empty()
            status_text.empty()
            st.error(f"❌ 审核失败: {str(e)}")
            return

    # ==================== 显示审核报告 ====================
    if st.session_state.audit_result:
        st.markdown("---")
        st.header("📊 审核报告")

        st.markdown(st.session_state.audit_result)

        with st.expander("📄 查看原始报告文本"):
            st.code(st.session_state.audit_result, language=None)

        st.markdown("---")
        col_dl1, col_dl2, col_dl3 = st.columns([1, 1, 1])
        with col_dl2:
            doc = audit_report_to_docx(st.session_state.audit_result)
            doc_bytes = io.BytesIO()
            doc.save(doc_bytes)
            doc_bytes.seek(0)

            st.download_button(
                label="📥 下载审核报告(Word)",
                data=doc_bytes,
                file_name="流程评审报告.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True,
                key="download_audit"
            )

    # ==================== 审核对话 ====================
    if st.session_state.audit_messages:
        st.markdown("---")
        st.header("💬 审核对话")

        chat_messages = st.session_state.audit_messages[3:] if len(st.session_state.audit_messages) > 3 else []

        for msg in chat_messages:
            with st.chat_message(msg["role"]):
                st.markdown(msg["content"])

        if prompt := st.chat_input("输入追问，如：'活动03的风险控制点是否充分？'", key="audit_chat"):
            st.session_state.audit_messages.append({"role": "user", "content": prompt})

            with st.chat_message("user"):
                st.markdown(prompt)

            with st.chat_message("assistant"):
                with st.spinner("思考中..."):
                    try:
                        client = DeepSeekClient(api_key=api_key, base_url=base_url, model=model)
                        response = client.chat(st.session_state.audit_messages)
                        st.markdown(response)
                        st.session_state.audit_messages.append({"role": "assistant", "content": response})
                    except Exception as e:
                        st.error(f"对话失败: {str(e)}")


# ==================== 主应用 ====================

def main():
    init_session_state()

    st.title("🏗️ 企业流程架构设计与审核智能体")
    st.markdown("**基于 APQC 国际标准的端到端三级流程架构设计与流程审核**")
    st.markdown("---")

    # 侧边栏 - API 配置
    with st.sidebar:
        st.header("⚙️ API 配置")

        def get_secret(key: str, default: str = "") -> str:
            try:
                if hasattr(st, 'secrets') and key in st.secrets:
                    return st.secrets[key]
            except Exception:
                pass
            return os.getenv(key, default)

        default_key = get_secret("DEEPSEEK_API_KEY", "")
        default_url = get_secret("DEEPSEEK_BASE_URL", DEEPSEEK_CONFIG["base_url"])
        default_model = get_secret("DEEPSEEK_MODEL", DEEPSEEK_CONFIG["model"])

        api_key = st.text_input(
            "DEEPSEEK_API_KEY",
            value=default_key,
            type="password",
            placeholder="sk-xxxxxxxx",
            help="输入你的 DeepSeek API Key"
        )

        base_url = st.text_input(
            "DEEPSEEK_BASE_URL",
            value=default_url,
            help="API 地址，一般无需修改"
        )

        model_options = ["deepseek-chat", "deepseek-coder"]
        if default_model and default_model not in model_options:
            model_options.insert(0, default_model)

        model = st.selectbox(
            "模型选择",
            options=model_options,
            index=model_options.index(default_model) if default_model in model_options else 0
        )

        st.markdown("---")

        if st.button("测试连接", key="test_conn"):
            if not api_key:
                st.error("请先输入 API Key")
            else:
                with st.spinner("测试中..."):
                    try:
                        _ = DeepSeekClient(api_key=api_key, base_url=base_url, model=model).test_connection()
                        st.success("✓ 连接成功")
                    except Exception as e:
                        st.error(f"✗ 连接失败: {str(e)}")

        st.markdown("---")
        st.markdown("""
        ### 使用说明

        1. 在上方输入 API Key
        2. 点击测试连接确认
        3. 选择「流程架构设计」或「流程审核」
        4. 输入信息并开始

        ### 支持文件
        .txt / .docx / .pdf
        """)

    # 主区域 - 功能选择
    tab_design, tab_audit = st.tabs(["🏗️ 流程架构设计", "🔍 流程审核"])

    with tab_design:
        render_design_page(api_key, base_url, model)

    with tab_audit:
        render_audit_page(api_key, base_url, model)


if __name__ == "__main__":
    main()
